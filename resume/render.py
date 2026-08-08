"""Renders data/resume_profile.yaml (+ optional per-job tailoring) into a DOCX resume.

Tailoring never invents facts — it only:
  - swaps in a JD-tailored summary paragraph (still truthful, written from resume_profile facts)
  - reorders bullets within a role so the most relevant ones lead
  - reorders skill names within a category so JD-relevant ones lead

Usage as a library:
    from resume.render import render
    render(resume_profile_dict, tailoring_dict, Path("out.docx"))

Usage standalone (renders the untailored base resume, useful for testing):
    python resume/render.py data/resume_profile.yaml out.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.enum.text import WD_LINE_SPACING

ROOT = Path(__file__).resolve().parent.parent

# Keeps a 6-YOE resume to one page: tight margins/fonts and a cap on how many
# bullets render per role (reordering already put the most relevant ones first,
# so truncating after reorder keeps the highest-signal bullets).
DEFAULT_MAX_BULLETS_PER_ROLE = 4


def _tighten(paragraph, space_after: int = 2, space_before: int = 0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _tighten_format(fmt, space_after: int = 2, space_before: int = 0) -> None:
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _add_hyperlink(paragraph, url: str, text: str, bold: bool = False, size_pt: float = 9.5) -> None:
    """Adds a clickable hyperlink run to a paragraph (python-docx has no
    built-in API for this; this is the standard OOXML-level workaround).
    Note: hyperlink runs live under w:hyperlink, not directly under w:p, so
    they don't show up in `paragraph.runs` — set size/bold here, not after."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    run_props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
    run_props.append(size)

    if bold:
        run_props.append(OxmlElement("w:b"))

    run_elem.append(run_props)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run_elem.append(text_elem)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _reorder(items: list, order: list[int] | None) -> list:
    if not order:
        return items
    if sorted(order) != list(range(len(items))):
        return items  # malformed order, fall back to original rather than dropping content
    return [items[i] for i in order]


def _reorder_by_name(names: list[str], highlight: list[str] | None) -> list[str]:
    if not highlight:
        return names
    highlight_lower = {h.lower() for h in highlight}
    front = [n for n in names if n.lower() in highlight_lower]
    back = [n for n in names if n.lower() not in highlight_lower]
    return front + back


def render(resume_profile: dict, tailoring: dict, output_path: Path) -> None:
    tailoring = tailoring or {}
    max_bullets = tailoring.get("max_bullets_per_role", DEFAULT_MAX_BULLETS_PER_ROLE)
    doc = Document()

    section = doc.sections[0]
    section.page_height, section.page_width = Cm(29.7), Cm(21.0)  # A4
    section.top_margin = section.bottom_margin = Cm(1.27)
    section.left_margin = section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    _tighten_format(style.paragraph_format)

    bullet_style = doc.styles["List Bullet"]
    bullet_style.font.name = "Calibri"
    bullet_style.font.size = Pt(9.5)
    _tighten_format(bullet_style.paragraph_format, space_after=1)

    contact = resume_profile["contact"]
    title = doc.add_heading(contact["name"], level=0)
    title.alignment = 1
    title.runs[0].font.size = Pt(18)
    _tighten(title, space_after=1)

    p = doc.add_paragraph()
    p.alignment = 1
    _tighten(p, space_after=6)
    parts = [("text", contact["phone"]), ("text", contact["email"])]
    if contact.get("linkedin"):
        parts.append(("link", contact["linkedin"]))
    if contact.get("github"):
        parts.append(("link", contact["github"]))
    for i, (kind, value) in enumerate(parts):
        if i > 0:
            p.add_run(" | ").font.size = Pt(9.5)
        if kind == "link":
            _add_hyperlink(p, value, value)
        else:
            p.add_run(value).font.size = Pt(9.5)

    def section_heading(text: str):
        h = doc.add_heading(text, level=1)
        h.runs[0].font.size = Pt(11.5)
        _tighten(h, space_before=6, space_after=2)
        return h

    section_heading("Professional Summary")
    _tighten(doc.add_paragraph(tailoring.get("summary") or resume_profile["summary_base"].strip()))

    section_heading("Technical Skills")
    skills = resume_profile["skills"]
    highlight = tailoring.get("skills_highlight", [])
    labels = {
        "languages": "Languages",
        "frameworks": "Frameworks",
        "distributed_systems": "Distributed Systems",
        "cloud_infra": "Cloud & Infrastructure",
        "databases": "Databases",
        "tooling": "Tooling & Practices",
        "ai": "AI",
    }
    for key, label in labels.items():
        values = skills.get(key, [])
        if not values:
            continue
        ordered = _reorder_by_name(values, highlight)
        _tighten(doc.add_paragraph(f"{label}: {', '.join(ordered)}"))

    section_heading("Experience")
    experience_order = tailoring.get("experience_order", {})
    for role in resume_profile["experience"]:
        heading = doc.add_paragraph()
        run = heading.add_run(f"{role['title']} — {role['company']} ({role['start']} to {role['end']})")
        run.bold = True
        _tighten(heading, space_before=4)

        bullets = [b["text"] for b in role["bullets"]]
        bullets = _reorder(bullets, experience_order.get(role["company"]))
        for bullet_text in bullets[:max_bullets]:
            doc.add_paragraph(bullet_text, style="List Bullet")

    if resume_profile.get("open_source"):
        section_heading("Open Source Contributions")
        for proj in resume_profile["open_source"]:
            p = doc.add_paragraph()
            _tighten(p)
            if proj.get("project_url"):
                _add_hyperlink(p, proj["project_url"], proj["project"], bold=True, size_pt=10)
            else:
                p.add_run(proj["project"]).bold = True
            p.add_run(f" — {proj['role']}").bold = True

            for bullet in proj["bullets"][:2]:
                # bullets may be a plain string or {text, links: [{label, url}]}
                text = bullet["text"] if isinstance(bullet, dict) else bullet
                links = bullet.get("links", []) if isinstance(bullet, dict) else []
                bp = doc.add_paragraph(text, style="List Bullet")
                for link in links:
                    bp.add_run(" ")
                    _add_hyperlink(bp, link["url"], link["label"], size_pt=9.5)

    if resume_profile.get("achievements"):
        section_heading("Achievements & Awards")
        for ach in resume_profile["achievements"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{ach['title']}: ")
            run.bold = True
            p.add_run(ach["text"])
            _tighten(p)

    section_heading("Education")
    for edu in resume_profile["education"]:
        _tighten(doc.add_paragraph(f"{edu['school']} — {edu['degree']} ({edu['start']}-{edu['end']})"))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    if len(sys.argv) != 3:
        print("Usage: python resume/render.py <resume_profile.yaml> <output.docx>")
        sys.exit(1)

    with open(sys.argv[1], "r", encoding="utf-8") as f:
        resume_profile = yaml.safe_load(f)

    render(resume_profile, {}, Path(sys.argv[2]))
    print(f"Wrote {sys.argv[2]}")


if __name__ == "__main__":
    main()
